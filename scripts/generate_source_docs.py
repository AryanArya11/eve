"""Generate beginner-friendly Word docs for each eve source module."""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

OUTPUT_DIR = Path(__file__).resolve().parent.parent / "docs" / "source-guides"

PROJECT_INTRO = (
    "Eve is a distributed batch-computing system for trusted local networks. "
    "It lets you submit a batch of small jobs (tasks), queue them, run them in "
    "separate worker processes, and collect structured success or failure results. "
    "The files in src/eve/ are the building blocks of that system."
)


def add_title(doc: Document, text: str) -> None:
    doc.add_heading(text, level=0)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def add_method_block(
    doc: Document,
    name: str,
    what: str,
    why: str,
    details: str | None = None,
) -> None:
    add_heading(doc, name, level=2)
    add_body(doc, f"What it does: {what}")
    add_body(doc, f"Why Eve needs it: {why}")
    if details:
        add_body(doc, details)


def write_attempt_doc() -> None:
    doc = Document()
    add_title(doc, "attempt.py — Execution Attempt Model")
    add_body(doc, PROJECT_INTRO)
    add_heading(doc, "What this file is for")
    add_body(
        doc,
        "When Eve runs a task, it might need more than one try (for example, after a "
        "network glitch or a retry policy). An ExecutionAttempt represents one specific "
        "try to run one task. It stores identity information only — not whether the "
        "run succeeded, which worker ran it, or when it finished. Those details live "
        "in other modules (state.py, outcome.py).",
    )
    add_heading(doc, "Class: ExecutionAttempt")
    add_body(
        doc,
        "A frozen (immutable) dataclass. Once created, its fields cannot change. "
        "That makes it safe to pass between the coordinator and worker processes "
        "without accidental edits.",
    )
    add_method_block(
        doc,
        "Fields: attempt_id, task_id, attempt_number",
        "attempt_id is a unique name for this try. task_id links the attempt to "
        "the Task it is trying to run. attempt_number is a 1-based counter (1 = first "
        "try, 2 = second try, and so on).",
        "Distributed systems need to distinguish 'task A, try 1' from 'task A, try 2' "
        "so retries, logging, and results do not get mixed up.",
    )
    add_method_block(
        doc,
        "__post_init__(self)",
        "Runs automatically right after the object is created. It calls the private "
        "validators on all three fields so bad data is rejected immediately.",
        "Eve fails fast at the boundary. Invalid attempts never enter the queue or "
        "reach a worker, which keeps the batch system predictable.",
    )
    add_method_block(
        doc,
        "_validate_field_name(field_name, value) [static]",
        "Checks that attempt_id and task_id are non-empty strings with no leading or "
        "trailing spaces. Raises TypeError for non-strings and ValueError for blank "
        "or padded IDs.",
        "Stable string IDs are how the coordinator, queue, and workers refer to the "
        "same logical unit of work across processes.",
    )
    add_method_block(
        doc,
        "_validate_attempt_number(value) [static]",
        "Ensures attempt_number is a real integer (not a bool), and is at least 1.",
        "Retry numbering must be unambiguous. Treating True/False as integers would "
        "be confusing; starting at 1 matches human-friendly 'first attempt' language.",
    )
    doc.save(OUTPUT_DIR / "attempt.docx")


def write_task_doc() -> None:
    doc = Document()
    add_title(doc, "task.py — Task Model")
    add_body(doc, PROJECT_INTRO)
    add_heading(doc, "What this file is for")
    add_body(
        doc,
        "A Task is the smallest unit of work Eve understands: 'run this named workload "
        "with this input data.' It describes what to do, not how execution is going "
        "or what happened afterward.",
    )
    add_heading(doc, "Class: Task")
    add_body(
        doc,
        "Frozen dataclass with task_id, workload, and payload.",
    )
    add_method_block(
        doc,
        "Fields: task_id, workload, payload",
        "task_id uniquely names this piece of work inside a job. workload is a string "
        "key such as 'sum-range' or 'echo' that selects a registered handler in "
        "workloads.py. payload is the input object passed to that handler (often a dict).",
        "Separating 'what to run' (Task) from 'one try at running it' (ExecutionAttempt) "
        "lets Eve retry tasks and track state without duplicating workload definitions.",
    )
    add_method_block(
        doc,
        "__post_init__(self)",
        "Validates task_id and workload using _validate_text_field.",
        "Only well-formed workload names reach the registry in workloads.py, where "
        "unknown names are rejected anyway — but validating early gives clearer errors.",
    )
    add_method_block(
        doc,
        "_validate_text_field(field_name, value) [static]",
        "Same identity rules as other Eve IDs: must be a non-blank string without "
        "surrounding whitespace.",
        "Consistent ID rules across the codebase make serialization and logging reliable.",
    )
    add_body(
        doc,
        "Note: payload is not validated here. Picklability and shape are checked when "
        "workloads run or when data crosses a process boundary.",
    )
    doc.save(OUTPUT_DIR / "task.docx")


def write_job_doc() -> None:
    doc = Document()
    add_title(doc, "job.py — Job Model")
    add_body(doc, PROJECT_INTRO)
    add_heading(doc, "What this file is for")
    add_body(
        doc,
        "Users often submit several related tasks at once (a batch). A Job groups "
        "those tasks under one job_id so the coordinator can treat them as a single "
        "submission.",
    )
    add_heading(doc, "Class: Job")
    add_method_block(
        doc,
        "Fields: job_id, tasks",
        "job_id names the batch. tasks is an immutable tuple of Task objects.",
        "Batch computing is Eve's core use case — jobs are the container for that batch.",
    )
    add_method_block(
        doc,
        "__post_init__(self)",
        "Validates job_id and the tasks tuple.",
        "A job with invalid or duplicate tasks would corrupt scheduling and result "
        "matching later.",
    )
    add_method_block(
        doc,
        "_validate_text_field(field_name, value) [static]",
        "Standard Eve ID validation for job_id.",
        "Jobs are referenced by ID across the system like tasks and attempts.",
    )
    add_method_block(
        doc,
        "_validate_instances(task) [static]",
        "Ensures tasks is a non-empty tuple where every element is a Task and every "
        "task_id is unique.",
        "Duplicate task IDs inside one job would make it impossible to know which "
        "result belongs to which task.",
    )
    add_body(
        doc,
        "Jobs do not store runtime state (queued, running, etc.). That is modeled "
        "separately in state.py.",
    )
    doc.save(OUTPUT_DIR / "job.docx")


def write_outcome_doc() -> None:
    doc = Document()
    add_title(doc, "outcome.py — Outcome Model")
    add_body(doc, PROJECT_INTRO)
    add_heading(doc, "What this file is for")
    add_body(
        doc,
        "After a worker runs a task, the coordinator needs a structured answer: "
        "did it work, what value came back, or what went wrong? This file defines "
        "those result types.",
    )
    add_heading(doc, "Enum: OutcomeStatus")
    add_method_block(
        doc,
        "SUCCESS and FAILURE",
        "Two allowed result states. SUCCESS means the workload completed without "
        "raising an exception. FAILURE means something went wrong during execution.",
        "A small, explicit enum keeps result handling simple for batch reporting "
        "and retries.",
    )
    add_heading(doc, "Class: TaskError")
    add_method_block(
        doc,
        "Fields: error_type, message",
        "Stores the exception class name (e.g. TypeError) and a human-readable message.",
        "Raw tracebacks are heavy for batch systems; structured errors are easier to "
        "aggregate, log, and show to users.",
    )
    add_method_block(
        doc,
        "__post_init__(self) and _validate_text_field",
        "Both error fields must be non-blank strings without extra whitespace.",
        "Empty error messages would make failures impossible to debug.",
    )
    add_heading(doc, "Class: TaskOutcome")
    add_method_block(
        doc,
        "Fields: task_id, status, value, error",
        "Links a result to a task_id. On success, value holds the return value and "
        "error is None. On failure, error is a TaskError and value must be None.",
        "This is the contract between workers and the coordinator: one outcome per "
        "completed attempt.",
    )
    add_method_block(
        doc,
        "__post_init__(self)",
        "Validates task_id, status, and the success/failure field combination.",
        "Prevents contradictory outcomes (e.g. SUCCESS with an error object) that "
        "would break downstream logic.",
    )
    add_method_block(
        doc,
        "_validate_status(value)",
        "Ensures status is an OutcomeStatus enum member, not a raw string.",
        "Type safety avoids subtle bugs when comparing results.",
    )
    add_method_block(
        doc,
        "_validate_outcome_combination(status, value, error)",
        "SUCCESS cannot have an error; FAILURE must have a TaskError and cannot "
        "carry a return value.",
        "Clear rules make it easy for the coordinator to branch on success vs failure.",
    )
    doc.save(OUTPUT_DIR / "outcome.docx")


def write_state_doc() -> None:
    doc = Document()
    add_title(doc, "state.py — Execution State Machine")
    add_body(doc, PROJECT_INTRO)
    add_heading(doc, "What this file is for")
    add_body(
        doc,
        "Each execution attempt moves through lifecycle stages: submitted, queued, "
        "running, and finally succeeded, failed, or cancelled. This file defines "
        "those states and which moves between them are legal.",
    )
    add_heading(doc, "Enum: ExecutionState")
    add_body(doc, "The seven states and what they mean:")
    add_bullet(doc, "SUBMITTED — coordinator accepted the attempt")
    add_bullet(doc, "QUEUED — waiting in line for a worker")
    add_bullet(doc, "ASSIGNED — a worker has been chosen")
    add_bullet(doc, "RUNNING — work is in progress")
    add_bullet(doc, "SUCCEEDED — finished successfully (terminal)")
    add_bullet(doc, "FAILED — finished with error (terminal)")
    add_bullet(doc, "CANCELLED — stopped on purpose (terminal)")
    add_heading(doc, "Constants")
    add_method_block(
        doc,
        "TERMINAL_STATES",
        "A frozenset of SUCCEEDED, FAILED, and CANCELLED.",
        "Terminal states mean 'no further progress' — important for knowing when "
        "an attempt is done.",
    )
    add_method_block(
        doc,
        "LEGAL_TRANSITIONS",
        "A dictionary mapping each state to the set of states you may move to next. "
        "For example, SUBMITTED may go to QUEUED or CANCELLED, but not directly to RUNNING.",
        "Prevents impossible lifecycles (like running before queued) in a distributed "
        "coordinator.",
    )
    add_method_block(
        doc,
        "is_terminal(state)",
        "Returns True if state is in TERMINAL_STATES. Raises TypeError if state is "
        "not an ExecutionState.",
        "Lets the coordinator quickly check whether to keep scheduling or close out "
        "an attempt.",
    )
    add_method_block(
        doc,
        "can_transition(current, target)",
        "Returns True if moving from current to target is allowed. Validates both "
        "arguments are ExecutionState values.",
        "Non-throwing check useful for UI or planning before committing a change.",
    )
    add_method_block(
        doc,
        "validate_transition(current, target)",
        "Calls can_transition; if illegal, raises ValueError with a clear message.",
        "Enforces rules at the moment the coordinator updates state, catching bugs early.",
    )
    doc.save(OUTPUT_DIR / "state.docx")


def write_local_queue_doc() -> None:
    doc = Document()
    add_title(doc, "local_queue.py — Local Execution Queue")
    add_body(doc, PROJECT_INTRO)
    add_heading(doc, "What this file is for")
    add_body(
        doc,
        "The coordinator needs a fair, ordered waiting line for work headed to workers. "
        "This module pairs a Task with its ExecutionAttempt and provides a FIFO queue.",
    )
    add_heading(doc, "Class: QueuedExecution")
    add_method_block(
        doc,
        "Fields: task, attempt",
        "Bundles the work definition (Task) with the specific try (ExecutionAttempt).",
        "Workers must run the right task and record results against the right attempt. "
        "Bundling prevents mismatched IDs.",
    )
    add_method_block(
        doc,
        "__post_init__(self)",
        "Validates types and that task.task_id == attempt.task_id.",
        "If IDs disagreed, Eve could execute task A but attribute the outcome to task B.",
    )
    add_method_block(
        doc,
        "_validate_task_type / _validate_attempt_type",
        "Ensure task is a Task and attempt is an ExecutionAttempt.",
        "Type checks at the queue boundary stop garbage data from entering the pipeline.",
    )
    add_method_block(
        doc,
        "_validate_matching_ids(task, attempt)",
        "Raises ValueError when task_id fields differ.",
        "Identity consistency is critical in batch systems with retries.",
    )
    add_heading(doc, "Class: LocalExecutionQueue")
    add_body(
        doc,
        "A mutable FIFO queue owned by the coordinator. Workers receive items after "
        "dequeue; they do not touch the queue directly.",
    )
    add_method_block(
        doc,
        "__init__(self)",
        "Creates an empty deque (items) and a set (queued_attempt_id) for duplicate detection.",
        "The set gives O(1) lookup so the same attempt cannot be queued twice at once.",
    )
    add_method_block(
        doc,
        "enqueue(item)",
        "Adds a QueuedExecution to the back. Rejects non-QueuedExecution items and "
        "duplicate attempt_id values.",
        "Duplicate attempts would cause double execution or ambiguous results.",
    )
    add_method_block(
        doc,
        "dequeue()",
        "Removes and returns the oldest item. Raises IndexError if empty. Removes "
        "the attempt_id from the tracking set.",
        "FIFO ordering means first-submitted work is handled first — fair batch scheduling.",
    )
    add_method_block(
        doc,
        "peek()",
        "Returns the next item without removing it. Raises IndexError if empty.",
        "Lets the coordinator inspect the head of the line without committing to assign it.",
    )
    add_method_block(
        doc,
        "is_empty()",
        "Returns True when there are no waiting items.",
        "Simple status check for scheduling loops.",
    )
    add_method_block(
        doc,
        "length()",
        "Returns the number of items waiting.",
        "Useful for metrics and backpressure decisions.",
    )
    doc.save(OUTPUT_DIR / "local_queue.docx")


def write_workloads_doc() -> None:
    doc = Document()
    add_title(doc, "workloads.py — Workload Registry and Handlers")
    add_body(doc, PROJECT_INTRO)
    add_heading(doc, "What this file is for")
    add_body(
        doc,
        "Tasks name a workload string, not arbitrary Python code. This file maps "
        "allowed names to real functions and runs them safely. It is the security "
        "boundary: only registered handlers execute.",
    )
    add_method_block(
        doc,
        "sum_range(payload)",
        "Expects a dict with integer start and stop (not bools). Sums integers from "
        "start inclusive to stop exclusive (like Python's range(start, stop)). "
        "Returns an int.",
        "A concrete example workload for testing batch math jobs and validating "
        "payload contracts.",
    )
    add_method_block(
        doc,
        "echo(payload)",
        "Returns the payload unchanged, including None.",
        "Useful for testing that a successful None result is not confused with a failure.",
    )
    add_method_block(
        doc,
        "_WORKLOAD_HANDLERS (private dict)",
        "Maps string names to handler functions, e.g. 'sum-range' -> sum_range.",
        "An allowlist: unknown workload strings raise ValueError instead of running "
        "arbitrary code (no eval).",
    )
    add_method_block(
        doc,
        "execute_workload(workload, payload)",
        "Looks up workload in the registry and calls the handler with payload. "
        "Raises TypeError if workload is not a string; ValueError if unknown.",
        "Single entry point for worker.py so all execution goes through the same "
        "safe dispatch logic.",
    )
    add_body(
        doc,
        "Handlers are module-level functions so they can be pickled for multiprocessing "
        "worker processes (especially on Windows with spawn).",
    )
    doc.save(OUTPUT_DIR / "workloads.docx")


def write_worker_doc() -> None:
    doc = Document()
    add_title(doc, "worker.py — Worker Execution Logic")
    add_body(doc, PROJECT_INTRO)
    add_heading(doc, "What this file is for")
    add_body(
        doc,
        "This is the function a worker process actually runs: take a QueuedExecution, "
        "run its workload, and return a TaskOutcome. It converts Python exceptions "
        "into structured failures.",
    )
    add_method_block(
        doc,
        "execute_queued_execution(execution)",
        "Validates execution is a QueuedExecution. Calls execute_workload with the "
        "task's workload and payload. On success, returns TaskOutcome with "
        "OutcomeStatus.SUCCESS and the return value. On any Exception, builds a "
        "TaskError (type name + message) and returns FAILURE with value None.",
        "Workers must never crash silently — the coordinator needs a TaskOutcome for "
        "every attempt so batch jobs can complete and report partial failures.",
    )
    add_body(
        doc,
        "Invalid input (wrong type) raises TypeError immediately instead of producing "
        "a TaskOutcome, because that is a programming error, not a workload failure.",
    )
    doc.save(OUTPUT_DIR / "worker.docx")


def write_process_executor_doc() -> None:
    doc = Document()
    add_title(doc, "process_executor.py — Process Boundary Execution")
    add_body(doc, PROJECT_INTRO)
    add_heading(doc, "What this file is for")
    add_body(
        doc,
        "Eve runs workloads outside the coordinator process for isolation. This module "
        "spawns a worker process, runs execute_queued_execution there, and wraps the "
        "result with metadata (attempt ID and worker PID).",
    )
    add_heading(doc, "Class: ProcessExecutionResult")
    add_method_block(
        doc,
        "Fields: attempt_id, worker_pid, outcome",
        "attempt_id ties the result to the try. worker_pid is the OS process ID of "
        "the worker. outcome is the TaskOutcome from worker logic.",
        "The coordinator can verify work really ran in another process, not in-process.",
    )
    add_method_block(
        doc,
        "__post_init__ and validators",
        "Validates attempt_id (text rules), worker_pid (positive int, not bool), "
        "and outcome (must be TaskOutcome).",
        "Structured results must be trustworthy before updating job state.",
    )
    add_method_block(
        doc,
        "_execute_with_process_metadata(execution) [internal]",
        "Calls execute_queued_execution, then packages ProcessExecutionResult with "
        "getpid() as worker_pid.",
        "Runs inside the child process; adds process identity to the return value.",
    )
    add_method_block(
        doc,
        "execute_in_process(execution)",
        "Public API: validates QueuedExecution, creates a ProcessPoolExecutor with "
        "max_workers=1 and multiprocessing 'spawn' context, submits "
        "_execute_with_process_metadata, waits for future.result(), returns "
        "ProcessExecutionResult.",
        "Spawn + one worker per call gives deterministic, isolated execution in v0.1. "
        "Infrastructure errors (process start, pickle failures) propagate to the "
        "coordinator; workload errors become FAILURE outcomes inside the worker.",
    )
    doc.save(OUTPUT_DIR / "process_executor.docx")


def write_init_doc() -> None:
    doc = Document()
    add_title(doc, "__init__.py — Package Entry Point")
    add_body(doc, PROJECT_INTRO)
    add_heading(doc, "What this file is for")
    add_body(
        doc,
        "In Python, a folder with __init__.py is treated as a package named eve. "
        "This file is currently empty, which is normal for early versions.",
    )
    add_heading(doc, "Why it exists")
    add_body(
        doc,
        "It marks src/eve as importable (from eve.task import Task, etc.). Later "
        "versions may re-export common types here so users can write "
        "import eve instead of many submodule imports.",
    )
    add_body(
        doc,
        "There are no classes or methods in this file yet — nothing to document "
        "method-by-method beyond its packaging role.",
    )
    doc.save(OUTPUT_DIR / "__init__.docx")


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    write_init_doc()
    write_attempt_doc()
    write_task_doc()
    write_job_doc()
    write_outcome_doc()
    write_state_doc()
    write_local_queue_doc()
    write_workloads_doc()
    write_worker_doc()
    write_process_executor_doc()
    print(f"Wrote Word docs to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
